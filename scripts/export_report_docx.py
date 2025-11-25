"""
Generate a Word (DOCX) report summarizing the Predicting Price Moves project.
The script pulls existing CSV summaries and figures produced during analysis.
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


def add_table_from_dataframe(doc: Document, df: pd.DataFrame, headers: list[str]) -> None:
    """Insert a styled table constructed from a dataframe."""
    table = doc.add_table(rows=len(df) + 1, cols=len(headers))
    table.style = "Light List Accent 1"
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row_idx, row in enumerate(df.itertuples(index=False), start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = str(value)


def add_picture_if_exists(doc: Document, path: Path, caption: str, width_in: float = 5.5) -> None:
    """Embed an image when the file exists."""
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        doc.add_paragraph(f"Figure: {caption}")


def main() -> Path:
    project_root = Path(".").resolve()
    output_dir = project_root / "output"
    figures_dir = output_dir / "report_figures"

    doc = Document()
    doc.add_heading("Predicting Price Moves with News Sentiment", 0)
    doc.add_paragraph(
        "Comprehensive multi-branch analysis for Nova Financial Solutions "
        "(Setup, Task-1, Task-2, Task-3). Generated directly from the repository snapshot.\n"
    )

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Nova Financial Solutions curates a large-scale Financial News and Stock Price Integration Dataset (FNSPID) "
        "to connect qualitative headlines with quantitative market moves. The repository combines data engineering "
        "modules, exploratory notebooks, technical analysis scripts, and sentiment/price correlation tooling. "
        "This report consolidates branch deliverables, exploratory findings, diagrams, and recommendations."
    )

    # 2. Business Objective
    doc.add_heading("2. Business Objective", level=1)
    doc.add_paragraph(
        "Deliver a reproducible workflow that ingests analyst/news headlines, enriches them with NLP-derived structure, "
        "couples the insights with equity OHLCV data, and primes downstream models to predict price moves driven by "
        "sentiment shocks."
    )
    goals = [
        "Standardize the environment setup path so analysts can recreate the toolchain quickly.",
        "Run exhaustive Exploratory Data Analysis (EDA) on the news corpus to surface publication behavior.",
        "Compute high-fidelity technical indicators for flagship tickers as modeling covariates.",
        "Prototype sentiment-versus-return correlation to validate that aggregated news tone is predictive.",
    ]
    for goal in goals:
        doc.add_paragraph(goal, style="List Bullet")

    # 3. Branch summaries with diagrams
    doc.add_heading("3. Independent Discussion of Each Git Branch", level=1)
    branch_items = [
        (
            "3.1 Setup Branch",
            "Defines the repository layout (data/, src/, scripts/, notebooks/, output/), dependency manifest "
            "(requirements.txt), and bootstrap instructions (INSTALLATION_INSTRUCTIONS.md, PACKAGE_INSTALLATION.md). "
            "GitHub Actions ensures dependencies at least install on CI, keeping onboarding friction low.",
            "setup_flow.png",
        ),
        (
            "3.2 Task-1 Branch (News EDA)",
            "Implements DataLoader, EDAAnalyzer, TopicModeler, and PublisherAnalyzer. scripts/run_eda.py orchestrates "
            "preprocessing, descriptive stats, temporal profiling, spike detection, LDA/BERTopic modeling, and "
            "publisher/topic exports saved under output/data/. Visual artifacts (histograms, heatmaps, topic charts) "
            "quantify publisher dominance and thematic concentration.",
            "task1_flow.png",
        ),
        (
            "3.3 Task-2 Branch (Technical Analysis)",
            "scripts/technical_analysis.py loads ticker OHLCV (e.g., AAPL.csv), cleans anomalies, applies TA-Lib "
            "indicators (SMA/EMA/RSI/MACD/Bollinger/ATR) and PyNance risk metrics (returns, volatility, Sharpe). "
            "It produces dashboards, a correlation matrix, and summary metrics JSON suitable for overlaying with "
            "sentiment features.",
            "task2_flow.png",
        ),
        (
            "3.4 Task-3 Branch (Sentiment vs Returns)",
            "scripts/news_sentiment_stock_correlation.py filters articles_with_topics.csv by ticker, scores TextBlob "
            "polarity per headline, aggregates by day, merges with daily returns, and reports Pearson correlation. "
            "Current coverage (≈10 articles per ticker concentrated in Jun‑2020) limits statistical power; the section "
            "documents this explicitly to set expectations for future scaling.",
            "task3_flow.png",
        ),
    ]
    for title, text, fig_name in branch_items:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
        add_picture_if_exists(doc, figures_dir / fig_name, fig_name)

    # 4. Methodology
    doc.add_heading("4. Methodology", level=1)
    methodology = [
        "Data engineering: schema validation, datetime parsing, headline length/word counts, publisher-domain normalization, NA culling.",
        "Exploratory analytics: descriptive stats, publisher rankings, temporal histograms, weekday/month time series, spike detection.",
        "Text mining: NLP preprocessing (nltk, spaCy), keyword extraction, LDA topic modeling (gensim) with optional BERTopic + pyLDAvis.",
        "Publisher intelligence: comparative keyword usage, topic preferences, domain statistics, visualization exports.",
        "Technical analysis: TA-Lib indicators and PyNance risk metrics (rolling volatility, cumulative returns, Sharpe, autocorrelation).",
        "Sentiment correlation: TextBlob polarity scoring per headline, daily aggregation, Pearson correlation with price pct-change.",
    ]
    for step in methodology:
        doc.add_paragraph(step, style="List Number")

    # 5. EDA section with tables and plots
    doc.add_heading("5. Exploratory Data Analysis (EDA)", level=1)
    summary_csv = output_dir / "report_summary_stats.csv"
    if summary_csv.exists():
        stats = pd.read_csv(summary_csv)
        stats.columns = ["metric", "value"]
        doc.add_paragraph("Dataset descriptors:")
        table = doc.add_table(rows=len(stats) + 1, cols=2)
        table.style = "Light List"
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        for idx, row in enumerate(stats.itertuples(index=False), start=1):
            metric_text = str(row.metric).replace("_", " ").title()
            table.cell(idx, 0).text = metric_text
            table.cell(idx, 1).text = str(row.value)

    dow_csv = output_dir / "report_dayofweek_counts.csv"
    if dow_csv.exists():
        doc.add_paragraph("\nPublication frequency by day of week:")
        dow_df = pd.read_csv(dow_csv)
        add_table_from_dataframe(doc, dow_df, ["day_of_week", "article_count"])

    for fig in [
        "articles_per_year.png",
        "monthly_trend.png",
        "headline_length_hist.png",
        "top_publishers.png",
        "topic_distribution.png",
    ]:
        add_picture_if_exists(doc, figures_dir / fig, fig)

    # 6. Diagrams (overall pipeline)
    doc.add_heading("6. Diagrams", level=1)
    add_picture_if_exists(doc, figures_dir / "project_pipeline.png", "project_pipeline.png")

    # 7. Findings
    doc.add_heading("7. Findings and Interpretation", level=1)
    findings = [
        (
            "Headline supply and publisher mix",
            "55,987 articles span Apr‑2011 → Jun‑2020, sourced from 225 publishers covering 6,204 tickers. "
            "Benzinga desks alone supply >55% of content (top-publisher table and bar chart), creating a single-source bias "
            "that should be weighted or normalized before modeling."
        ),
        (
            "Temporal concentration",
            "EDA plots show Thursday as the busiest news day (12,712 items) with pronounced spikes during Mar‑Jun 2020. "
            "Weekend coverage is negligible (<1%), implying that intraday strategies should focus on market hours while "
            "swing strategies should treat weekend sentiment as noise."
        ),
        (
            "Topic structure",
            "LDA topic distribution indicates Topic 3 (earnings beats/surprises) dominates with 13,257 articles, followed by ETF rotation "
            "and target-adjustment themes. This validates the feasibility of topic-aware sentiment factors (e.g., overweighting earnings "
            "language near reporting seasons)."
        ),
        (
            "Technical context",
            "Task-2 metrics for AAPL report annualized volatility ≈12.7%, Sharpe ≈1.07, and negligible lagged autocorrelation, indicating "
            "returns are largely driven by exogenous shocks. This strengthens the case for exogenous sentiment inputs."
        ),
        (
            "Correlation evidence and limitations",
            "The proof-of-concept Pearson analysis on AAPL yields −1.0 but is based on only two overlapping dates (9–10 Jun 2020). "
            "Such a tiny sample makes the coefficient statistically meaningless; the report therefore treats it as a placeholder and "
            "recommends collecting a broader rolling window (e.g., 12–18 months per ticker) before drawing inference."
        ),
    ]
    for title, detail in findings:
        doc.add_paragraph(f"{title}: {detail}", style="List Bullet")

    doc.add_heading("8. Investment Strategy Concepts (Hypothetical)", level=1)
    strategies = [
        "Earnings Momentum Overlay: When topic and keyword classifiers flag earnings-beat language combined with positive sentiment, "
        "tilt long exposures in the corresponding ticker or sector ETF for the subsequent 1–3 sessions; conversely, fade negative "
        "earnings sentiment signals during peak season.",
        "Publisher Quality Spread: Weight signals by historical accuracy of specific Benzinga desks or other publishers. "
        "A high-sentiment reading from a historically impactful source can trigger call-spread entries, while low sentiment could "
        "justify protective puts.",
        "Macro Shock Radar: Utilize spike-detection outputs to identify systemic news bursts (e.g., pandemic headlines). "
        "Combine with volatility metrics to size hedges (e.g., add VIX calls or sector-rotation trades) when sentiment spikes "
        "precede volatility expansion.",
        "Event-Driven Pair Trades: When topic modeling surfaces sector-specific narratives (FDA approvals, M&A), pair the affected "
        "ticker against its ETF to exploit sentiment-driven divergences over a short horizon.",
    ]
    for idea in strategies:
        doc.add_paragraph(idea, style="List Bullet")

    # 9. Challenges
    doc.add_heading("9. Challenges and Limitations", level=1)
    challenges = [
        "financial_news.csv is absent from version control, so analysts must supply it manually before running Task-1.",
        "Ticker coverage in articles_with_topics.csv is ~10 samples per symbol, limiting sentiment statistics.",
        "Heavy dependencies (TA-Lib, BERTopic, PyTorch) complicate environment setup without OS-specific guidance.",
        "CI pipeline only validates installation; automated tests/linting are missing.",
        "Scripts rely on sys.path manipulation instead of packaging src/ as an installable module.",
    ]
    for challenge in challenges:
        doc.add_paragraph(challenge, style="List Bullet")

    # 10. Conclusion and Recommendations
    doc.add_heading("10. Conclusion and Recommendations", level=1)
    doc.add_paragraph(
        "Expand the news ingestion window to provide deeper ticker histories, adopt finance-aware sentiment models (FinBERT/VADER), "
        "package src/ as an installable library, and enhance CI with pytest plus data validation. "
        "From a trading perspective, prioritize building rolling sentiment indices per topic and publisher, calibrate their "
        "lead/lag against technical indicators, and backtest the strategy templates above (earnings overlay, publisher quality spread, "
        "macro shock radar, pair trades) once sufficient history exists. Persist aggregated features (topics, publisher metrics, technical "
        "indicators, daily sentiment) in columnar storage to feed portfolio construction and risk systems."
    )

    output_path = output_dir / "Predicting_Price_Moves_Report.docx"
    try:
        doc.save(output_path)
        return output_path
    except PermissionError:
        fallback = output_dir / "Predicting_Price_Moves_Report_v2.docx"
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    saved_path = main()
    print(f"Word report saved to {saved_path}")

